"""Deterministic, local text extraction for competition documents.

No LLM or remote service is used here. The output preserves filenames and real
PDF page numbers so later chunks and citations can inherit trustworthy metadata.
"""

from __future__ import annotations

from pathlib import Path
from typing import Callable

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models import ExtractedDocument, ExtractedPage
from app.pdf_glyphs import repair_thai_glyphs


class ExtractionError(RuntimeError):
    """Base class for expected document extraction failures."""


class UnsupportedDocumentError(ExtractionError):
    """Raised when a file extension is not supported."""


class EmptyDocumentError(ExtractionError):
    """Raised when a document contains no extractable text."""


class InvalidDocumentError(ExtractionError):
    """Raised when a supported document cannot be parsed."""


class DocumentTextExtractor:
    """Extract PDF, TXT, and DOCX files into metadata-preserving source units."""

    SUPPORTED_EXTENSIONS = frozenset({".pdf", ".txt", ".docx"})

    def extract(self, path: str | Path) -> ExtractedDocument:
        source_path = Path(path).expanduser().resolve()
        if not source_path.exists():
            raise ExtractionError(f"Document does not exist: {source_path}")
        if not source_path.is_file():
            raise ExtractionError(f"Document path is not a file: {source_path}")

        handlers: dict[str, Callable[[Path], tuple[ExtractedPage, ...]]] = {
            ".pdf": self._extract_pdf,
            ".txt": self._extract_txt,
            ".docx": self._extract_docx,
        }
        suffix = source_path.suffix.lower()
        handler = handlers.get(suffix)
        if handler is None:
            supported = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise UnsupportedDocumentError(
                f"Unsupported file type {suffix or '(none)'}. Supported: {supported}"
            )

        pages = handler(source_path)
        if not pages or not any(page.text.strip() for page in pages):
            extra = " Scanned PDFs require a separate OCR step." if suffix == ".pdf" else ""
            raise EmptyDocumentError(
                f"No extractable text found in {source_path.name}.{extra}"
            )

        return ExtractedDocument(
            source_path=source_path,
            document=source_path.name,
            pages=pages,
        )

    @staticmethod
    def _clean_text(text: str) -> str:
        lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]
        return "\n".join(lines).strip()

    def _extract_pdf(self, path: Path) -> tuple[ExtractedPage, ...]:
        try:
            reader = PdfReader(path)
            repair_thai_glyphs(reader)
            pages = []
            for page_number, pdf_page in enumerate(reader.pages, start=1):
                text = self._clean_text(pdf_page.extract_text() or "")
                if text:
                    pages.append(
                        ExtractedPage(
                            document=path.name,
                            page=page_number,
                            text=text,
                        )
                    )
            return tuple(pages)
        except ExtractionError:
            raise
        except Exception as exc:
            raise InvalidDocumentError(
                f"Could not read PDF {path.name}: {exc}"
            ) from exc

    def _extract_txt(self, path: Path) -> tuple[ExtractedPage, ...]:
        text: str | None = None
        decode_errors: list[str] = []
        for encoding in ("utf-8-sig", "utf-8", "cp874"):
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError as exc:
                decode_errors.append(f"{encoding}: {exc}")
            except OSError as exc:
                raise InvalidDocumentError(
                    f"Could not read text file {path.name}: {exc}"
                ) from exc

        if text is None:
            raise InvalidDocumentError(
                f"Could not decode {path.name} as UTF-8 or Thai CP874. "
                + " | ".join(decode_errors)
            )

        cleaned = self._clean_text(text)
        return (ExtractedPage(document=path.name, page=None, text=cleaned),)

    def _extract_docx(self, path: Path) -> tuple[ExtractedPage, ...]:
        try:
            document = DocxDocument(path)
            blocks: list[str] = []

            blocks.extend(
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            )

            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(cell for cell in cells if cell)
                    if row_text:
                        blocks.append(row_text)

            text = self._clean_text("\n".join(blocks))
            return (ExtractedPage(document=path.name, page=None, text=text),)
        except Exception as exc:
            raise InvalidDocumentError(
                f"Could not read DOCX {path.name}: {exc}"
            ) from exc
