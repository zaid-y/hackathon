"""
Document Extraction Module for ThaiLLM RAG
Supports loading and chunking documents from PDF, DOCX, TXT, HTML, and web URLs
with Thai-aware text segmentation.
"""
import os
import re
import requests
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Callable
from dataclasses import dataclass, field
from urllib.parse import urlparse

# Optional imports with graceful fallback
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

# Thai tokenization (same pattern as retriever.py)
try:
    from pythainlp.tokenize import word_tokenize
    HAS_PYTHAINLP = True
except ImportError:
    HAS_PYTHAINLP = False

try:
    import deepcut
    HAS_DEEPCUT = True
except ImportError:
    HAS_DEEPCUT = False


@dataclass
class ExtractedDocument:
    """Document extracted from a source file/URL before chunking"""
    content: str
    metadata: Dict[str, Any]
    source_path: str
    page_number: Optional[int] = None


@dataclass
class ChunkingConfig:
    """Configuration for text chunking"""
    chunk_size: int = 500          # Target tokens per chunk
    chunk_overlap: int = 50        # Overlap tokens between chunks
    min_chunk_size: int = 50       # Minimum tokens for a valid chunk
    preserve_sentences: bool = True # Try to keep sentences intact
    preserve_paragraphs: bool = True  # Try to keep paragraphs intact
    adaptive_chunking: bool = True    # Use heading-aware chunking


class ThaiTokenizer:
    """
    Thai-aware tokenizer with multiple backend support.
    Priority: pythainlp (newmm) > deepcut > regex fallback
    """

    def __init__(self):
        self._backend = self._detect_backend()

    def _detect_backend(self) -> str:
        if HAS_PYTHAINLP:
            return "pythainlp"
        elif HAS_DEEPCUT:
            return "deepcut"
        return "regex"

    def tokenize(self, text: str) -> List[str]:
        """Tokenize text into words/tokens"""
        if self._backend == "pythainlp":
            try:
                tokens = word_tokenize(text, engine="newmm", keep_whitespace=False)
                return [t for t in tokens if t.strip()]
            except Exception:
                pass

        if self._backend == "deepcut" and HAS_DEEPCUT:
            try:
                tokens = deepcut.tokenize(text)
                return [t for t in tokens if t.strip()]
            except Exception:
                pass

        # Regex fallback: Thai characters + alphanumeric
        tokens = re.findall(r'[฀-๿]+|[a-zA-Z0-9]+', text)
        return [t for t in tokens if t.strip()]

    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        return len(self.tokenize(text))


class ThaiTextChunker:
    """
    Intelligent text chunker for Thai language.
    Respects Thai word boundaries, preserves sentence/paragraph structure,
    and supports heading-aware adaptive chunking.
    """

    # Thai heading patterns
    HEADING_PATTERNS = [
        r'^#{1,6}\s+.+$',              # Markdown headings
        r'^.{1,100}\n[-=]{3,}$',       # Underlined headings
        r'^\d+[.)]\s+.+$',             # Numbered list headings
        r'^[➤▶◆●■►▪]\s*.+$',           # Bullet headings
        r'^[ก-ฮ]\.\s+.+$',             # Thai numbered (ก. ข. ค.)
        r'^\([ก-ฮ]\)\s+.+$',           # Thai parenthesized (ก) (ข)
        r'^หัวข้อ\s*\d+',              # "หัวข้อ 1"
        r'^บทที่\s*\d+',               # "บทที่ 1"
        r'^ส่วนที่\s*\d+',             # "ส่วนที่ 1"
        r'^มาตรา\s*\d+',              # "มาตรา 1"
        r'^ข้อ\s*\d+',                # "ข้อ 1"
    ]

    def __init__(self, config: Optional[ChunkingConfig] = None):
        self.config = config or ChunkingConfig()
        self.tokenizer = ThaiTokenizer()
        self._heading_regex = re.compile('|'.join(self.HEADING_PATTERNS), re.MULTILINE)

    def chunk(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[ExtractedDocument]:
        """
        Split text into overlapping chunks respecting Thai word boundaries.

        Args:
            text: Input text to chunk
            metadata: Base metadata to attach to each chunk

        Returns:
            List of ExtractedDocument chunks
        """
        if not text or not text.strip():
            return []

        tokens = self.tokenizer.tokenize(text)
        if len(tokens) <= self.config.chunk_size:
            meta = (metadata or {}).copy()
            meta["chunk_index"] = 0
            meta["chunk_tokens"] = len(tokens)
            return [ExtractedDocument(
                content=text.strip(),
                metadata=meta,
                source_path=metadata.get("source_path", "") if metadata else ""
            )]

        base_metadata = metadata or {}
        chunks = []

        # Use adaptive (heading-aware) chunking if enabled
        if self.config.adaptive_chunking:
            chunks = self._adaptive_chunk(text, base_metadata)
        elif self.config.preserve_sentences:
            sentences = self._split_sentences(text)
            chunks = self._chunk_by_sentences(sentences, base_metadata)
        else:
            chunks = self._chunk_by_tokens(tokens, text, base_metadata)

        # Filter too-small chunks
        filtered = [c for c in chunks if self.tokenizer.count_tokens(c.content) >= self.config.min_chunk_size]
        return filtered

    def _adaptive_chunk(self, text: str, base_metadata: Dict) -> List[ExtractedDocument]:
        """
        Heading-aware chunking: split by headings, then chunk each section.
        This preserves document structure better than flat chunking.
        """
        # Find all headings with their positions
        headings = list(self._heading_regex.finditer(text))

        if not headings:
            # No headings found, fall back to sentence chunking
            sentences = self._split_sentences(text)
            return self._chunk_by_sentences(sentences, base_metadata)

        chunks = []
        chunk_index = 0
        prev_end = 0

        # Process each section between headings
        for i, heading_match in enumerate(headings):
            heading_text = heading_match.group()
            heading_start = heading_match.start()
            heading_end = heading_match.end()

            # Determine section boundaries
            section_start = heading_start
            section_end = headings[i + 1].start() if i + 1 < len(headings) else len(text)

            # Extract section content (including the heading)
            section_content = text[section_start:section_end].strip()

            if not section_content:
                continue

            # Extract heading level for metadata
            heading_level = self._get_heading_level(heading_text)

            # Chunk this section
            section_chunks = self._chunk_section(
                section_content,
                base_metadata,
                chunk_index,
                heading_text=heading_text,
                heading_level=heading_level
            )
            chunks.extend(section_chunks)
            chunk_index += len(section_chunks)

        return chunks

    def _get_heading_level(self, heading: str) -> int:
        """Determine heading level from heading text"""
        if heading.startswith('#'):
            return len(heading) - len(heading.lstrip('#'))
        if re.match(r'^.{1,100}\n[-=]{3,}$', heading):
            return 1 if '=' in heading else 2
        if re.match(r'^\d+[.)]\s+', heading):
            return 2
        if re.match(r'^[➤▶◆●■►▪]\s*', heading):
            return 3
        if re.match(r'^[ก-ฮ]\.\s+', heading) or re.match(r'^\([ก-ฮ]\)\s+', heading):
            return 2
        if any(prefix in heading for prefix in ['หัวข้อ', 'บทที่', 'ส่วนที่', 'มาตรา', 'ข้อ']):
            return 2
        return 3

    def _chunk_section(
        self,
        section_content: str,
        base_metadata: Dict,
        start_index: int,
        heading_text: str = "",
        heading_level: int = 0
    ) -> List[ExtractedDocument]:
        """Chunk a single section, preserving the heading in the first chunk"""
        # Split section into sentences
        sentences = self._split_sentences(section_content)
        if not sentences:
            return []

        chunks = []
        current_chunk = ""
        current_tokens = 0
        chunk_index = start_index
        first_chunk = True

        for sent in sentences:
            sent_tokens = self.tokenizer.count_tokens(sent)

            # If single sentence exceeds chunk_size, split it
            if sent_tokens > self.config.chunk_size:
                if current_chunk:
                    chunks.append(self._make_chunk(
                        current_chunk, base_metadata, chunk_index,
                        heading_text if first_chunk else "",
                        heading_level if first_chunk else 0
                    ))
                    chunk_index += 1
                    current_chunk = ""
                    current_tokens = 0
                    first_chunk = False

                # Split long sentence by tokens with overlap
                sub_chunks = self._split_long_sentence(
                    sent, base_metadata, chunk_index,
                    heading_text if first_chunk else "",
                    heading_level if first_chunk else 0
                )
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)
                first_chunk = False
                continue

            # Check if adding this sentence would exceed chunk_size
            if current_tokens + sent_tokens > self.config.chunk_size and current_chunk:
                chunks.append(self._make_chunk(
                    current_chunk, base_metadata, chunk_index,
                    heading_text if first_chunk else "",
                    heading_level if first_chunk else 0
                ))
                chunk_index += 1
                first_chunk = False

                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sent
                current_tokens = self.tokenizer.count_tokens(current_chunk)
            else:
                current_chunk = (current_chunk + " " + sent).strip() if current_chunk else sent
                current_tokens += sent_tokens

        # Add final chunk
        if current_chunk:
            chunks.append(self._make_chunk(
                current_chunk, base_metadata, chunk_index,
                heading_text if first_chunk else "",
                heading_level if first_chunk else 0
            ))

        return chunks

    def _split_long_sentence(
        self,
        sentence: str,
        base_metadata: Dict,
        start_index: int,
        heading_text: str = "",
        heading_level: int = 0
    ) -> List[ExtractedDocument]:
        """Split a sentence that's longer than chunk_size by tokens"""
        tokens = self.tokenizer.tokenize(sentence)
        chunks = []

        for i in range(0, len(tokens), self.config.chunk_size - self.config.chunk_overlap):
            chunk_tokens = tokens[i:i + self.config.chunk_size]
            # For Thai, join without spaces; for mixed, we need smarter joining
            chunk_text = self._join_tokens_smart(chunk_tokens)
            chunks.append(self._make_chunk(
                chunk_text, base_metadata, start_index + len(chunks),
                heading_text if len(chunks) == 0 else "",
                heading_level if len(chunks) == 0 else 0
            ))

        return chunks

    def _join_tokens_smart(self, tokens: List[str]) -> str:
        """Join tokens with appropriate spacing for Thai/English mixed text"""
        if not tokens:
            return ""

        result = tokens[0]
        for i in range(1, len(tokens)):
            prev = tokens[i-1]
            curr = tokens[i]

            # Add space if both are English/alphanumeric
            if re.match(r'^[a-zA-Z0-9]+$', prev) and re.match(r'^[a-zA-Z0-9]+$', curr):
                result += " " + curr
            # Add space if previous ends with English and current starts with English
            elif re.search(r'[a-zA-Z0-9]$', prev) and re.match(r'^[a-zA-Z0-9]', curr):
                result += " " + curr
            else:
                # Thai or mixed - no space
                result += curr

        return result

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences (Thai + English) with better handling"""
        # More comprehensive Thai/English sentence endings
        sentence_endings = r'[.!?。？！ฯ]+'

        # Split but keep the delimiter
        parts = re.split(f'({sentence_endings})', text)

        sentences = []
        current = ""
        for part in parts:
            current += part
            if re.match(sentence_endings, part):
                if current.strip():
                    sentences.append(current.strip())
                current = ""

        if current.strip():
            sentences.append(current.strip())

        # Merge very short sentences (likely abbreviations)
        merged = []
        for sent in sentences:
            if merged and len(sent) < 10 and not re.search(r'[.!?]$', merged[-1]):
                merged[-1] += " " + sent
            else:
                merged.append(sent)

        return merged if merged else [text]

    def _chunk_by_sentences(self, sentences: List[str], base_metadata: Dict) -> List[ExtractedDocument]:
        """Chunk by combining sentences up to chunk_size tokens"""
        chunks = []
        current_chunk = ""
        current_tokens = 0
        chunk_index = 0

        for sent in sentences:
            sent_tokens = self.tokenizer.count_tokens(sent)

            # If single sentence exceeds chunk_size, split it by tokens
            if sent_tokens > self.config.chunk_size:
                if current_chunk:
                    chunks.append(self._make_chunk(current_chunk, base_metadata, chunk_index))
                    chunk_index += 1
                    current_chunk = ""
                    current_tokens = 0

                # Split long sentence by tokens
                sub_chunks = self._split_long_sentence(sent, base_metadata, chunk_index)
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)
                continue

            # Check if adding this sentence would exceed chunk_size
            if current_tokens + sent_tokens > self.config.chunk_size and current_chunk:
                chunks.append(self._make_chunk(current_chunk, base_metadata, chunk_index))
                chunk_index += 1

                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sent
                current_tokens = self.tokenizer.count_tokens(current_chunk)
            else:
                current_chunk = (current_chunk + " " + sent).strip() if current_chunk else sent
                current_tokens += sent_tokens

        # Add final chunk
        if current_chunk:
            chunks.append(self._make_chunk(current_chunk, base_metadata, chunk_index))

        return chunks

    def _chunk_by_tokens(self, tokens: List[str], original_text: str, base_metadata: Dict) -> List[ExtractedDocument]:
        """Fallback: chunk purely by token count with overlap"""
        chunks = []
        chunk_index = 0

        for i in range(0, len(tokens), self.config.chunk_size - self.config.chunk_overlap):
            chunk_tokens = tokens[i:i + self.config.chunk_size]
            chunk_text = self._join_tokens_smart(chunk_tokens)
            chunks.append(self._make_chunk(chunk_text, base_metadata, chunk_index))
            chunk_index += 1

        return chunks

    def _get_overlap_text(self, text: str) -> str:
        """Get last N tokens for overlap, joined smartly"""
        tokens = self.tokenizer.tokenize(text)
        overlap_tokens = tokens[-self.config.chunk_overlap:] if len(tokens) > self.config.chunk_overlap else tokens
        return self._join_tokens_smart(overlap_tokens)

    def _make_chunk(
        self,
        content: str,
        base_metadata: Dict,
        chunk_index: int,
        heading_text: str = "",
        heading_level: int = 0
    ) -> ExtractedDocument:
        """Create ExtractedDocument with chunk metadata including heading info"""
        meta = base_metadata.copy()
        meta["chunk_index"] = chunk_index
        meta["chunk_tokens"] = self.tokenizer.count_tokens(content)

        # Add heading information for context
        if heading_text:
            meta["heading"] = heading_text.strip()
            meta["heading_level"] = heading_level

        return ExtractedDocument(
            content=content.strip(),
            metadata=meta,
            source_path=base_metadata.get("source_path", "")
        )


class DocumentExtractor(ABC):
    """Abstract base class for document extractors"""

    @abstractmethod
    def extract(self, source: str) -> List[ExtractedDocument]:
        """Extract documents from source (file path or URL)"""
        pass

    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions"""
        pass

    def _create_base_metadata(self, source: str, **extra) -> Dict[str, Any]:
        """Create base metadata dictionary"""
        return {
            "source_path": source,
            "source_name": os.path.basename(source) if not source.startswith(("http://", "https://")) else source,
            "extractor": self.__class__.__name__,
            **extra
        }


class PDFExtractor(DocumentExtractor):
    """Extract text from PDF files using pdfplumber or PyMuPDF"""

    def supported_extensions(self) -> List[str]:
        return [".pdf"]

    def extract(self, source: str) -> List[ExtractedDocument]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"PDF file not found: {source}")

        documents = []

        # Try pdfplumber first (better text extraction)
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(source) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            meta = self._create_base_metadata(source, page_number=page_num, total_pages=len(pdf.pages))
                            documents.append(ExtractedDocument(
                                content=text.strip(),
                                metadata=meta,
                                source_path=source,
                                page_number=page_num
                            ))
                return documents
            except Exception as e:
                print(f"[PDFExtractor] pdfplumber failed: {e}, trying PyMuPDF")

        # Fallback to PyMuPDF
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(source)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    if text and text.strip():
                        meta = self._create_base_metadata(source, page_number=page_num + 1, total_pages=len(doc))
                        documents.append(ExtractedDocument(
                            content=text.strip(),
                            metadata=meta,
                            source_path=source,
                            page_number=page_num + 1
                        ))
                doc.close()
                return documents
            except Exception as e:
                raise RuntimeError(f"PyMuPDF extraction failed: {e}")

        raise RuntimeError("No PDF extraction library available. Install pdfplumber or pymupdf.")


class DocxExtractor(DocumentExtractor):
    """Extract text from DOCX files"""

    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def extract(self, source: str) -> List[ExtractedDocument]:
        if not HAS_DOCX:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

        if not os.path.exists(source):
            raise FileNotFoundError(f"DOCX file not found: {source}")

        doc = DocxDocument(source)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)

        content = "\n\n".join(full_text)
        if not content:
            return []

        meta = self._create_base_metadata(source, total_paragraphs=len(doc.paragraphs))
        return [ExtractedDocument(
            content=content,
            metadata=meta,
            source_path=source
        )]


class TextExtractor(DocumentExtractor):
    """Extract text from plain text files (TXT, MD, etc.)"""

    def supported_extensions(self) -> List[str]:
        return [".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".xml", ".yaml", ".yml"]

    def extract(self, source: str) -> List[ExtractedDocument]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"Text file not found: {source}")

        with open(source, "r", encoding="utf-8") as f:
            content = f.read()

        if not content.strip():
            return []

        meta = self._create_base_metadata(source, file_size=len(content))
        return [ExtractedDocument(
            content=content,
            metadata=meta,
            source_path=source
        )]


class HTMLExtractor(DocumentExtractor):
    """Extract text from HTML files"""

    def supported_extensions(self) -> List[str]:
        return [".html", ".htm", ".xhtml"]

    def extract(self, source: str) -> List[ExtractedDocument]:
        if not HAS_BS4:
            raise RuntimeError("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")

        if not os.path.exists(source):
            raise FileNotFoundError(f"HTML file not found: {source}")

        with open(source, "r", encoding="utf-8") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator="\n", strip=True)

        # Clean up excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)

        if not text.strip():
            return []

        meta = self._create_base_metadata(source, title=soup.title.string if soup.title else None)
        return [ExtractedDocument(
            content=text.strip(),
            metadata=meta,
            source_path=source
        )]


class WebExtractor(DocumentExtractor):
    """Extract text from web URLs"""

    def supported_extensions(self) -> List[str]:
        return []  # URLs don't have extensions

    def supports_url(self, url: str) -> bool:
        """Check if this is a valid HTTP/HTTPS URL"""
        try:
            result = urlparse(url)
            return result.scheme in ("http", "https")
        except Exception:
            return False

    def extract(self, source: str) -> List[ExtractedDocument]:
        if not self.supports_url(source):
            raise ValueError(f"Invalid URL: {source}")

        if not HAS_BS4:
            raise RuntimeError("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")

        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            }
            response = requests.get(source, headers=headers, timeout=30)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, "html.parser")

            # Remove unwanted elements
            for elem in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
                elem.decompose()

            # Try to find main content
            main_content = soup.find("main") or soup.find("article") or soup.find("div", class_=re.compile(r"content|main|article", re.I))
            if main_content:
                text = main_content.get_text(separator="\n", strip=True)
            else:
                text = soup.get_text(separator="\n", strip=True)

            # Clean up
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r'[ \t]{2,}', ' ', text)

            if not text.strip():
                return []

            meta = self._create_base_metadata(
                source,
                title=soup.title.string if soup.title else None,
                url=source,
                status_code=response.status_code
            )
            return [ExtractedDocument(
                content=text.strip(),
                metadata=meta,
                source_path=source
            )]

        except requests.RequestException as e:
            raise RuntimeError(f"Failed to fetch URL: {e}")


# Factory functions

def create_extractor(source: str) -> DocumentExtractor:
    """
    Factory function to create appropriate extractor based on source.

    Args:
        source: File path or URL

    Returns:
        DocumentExtractor instance
    """
    # Check if URL
    if source.startswith(("http://", "https://")):
        return WebExtractor()

    # Get extension
    ext = os.path.splitext(source)[1].lower()

    # Map extensions to extractors
    if ext in [".pdf"]:
        return PDFExtractor()
    elif ext in [".docx", ".doc"]:
        return DocxExtractor()
    elif ext in [".html", ".htm", ".xhtml"]:
        return HTMLExtractor()
    elif ext in [".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".xml", ".yaml", ".yml"]:
        return TextExtractor()
    else:
        # Default to text extractor
        return TextExtractor()


def extract_documents(
    source: str,
    chunker: Optional[ThaiTextChunker] = None
) -> List[ExtractedDocument]:
    """
    High-level function to extract and chunk documents from a source.

    Args:
        source: File path or URL
        chunker: Optional ThaiTextChunker instance (creates default if not provided)

    Returns:
        List of chunked ExtractedDocument objects ready for RAG pipeline
    """
    extractor = create_extractor(source)
    raw_docs = extractor.extract(source)

    if not raw_docs:
        return []

    chunker = chunker or ThaiTextChunker()
    all_chunks = []

    for doc in raw_docs:
        chunks = chunker.chunk(doc.content, doc.metadata)
        all_chunks.extend(chunks)

    return all_chunks


def extract_multiple_sources(
    sources: List[str],
    chunker: Optional[ThaiTextChunker] = None
) -> List[ExtractedDocument]:
    """
    Extract and chunk documents from multiple sources.

    Args:
        sources: List of file paths or URLs
        chunker: Optional ThaiTextChunker instance

    Returns:
        Combined list of all chunked documents
    """
    all_chunks = []
    for source in sources:
        try:
            chunks = extract_documents(source, chunker)
            all_chunks.extend(chunks)
        except Exception as e:
            print(f"[extract_multiple_sources] Failed to process {source}: {e}")
            continue
    return all_chunks


# Convenience function for RAG pipeline integration
def load_documents_for_rag(
    sources: List[str],
    chunk_size: int = 500,
    chunk_overlap: int = 50
) -> List[Dict[str, Any]]:
    """
    Load documents from sources and format for RAG pipeline's add_documents().

    Args:
        sources: List of file paths or URLs
        chunk_size: Target tokens per chunk
        chunk_overlap: Overlap tokens between chunks

    Returns:
        List of dicts with 'content' and 'metadata' keys for pipeline.add_texts()
    """
    chunker = ThaiTextChunker(ChunkingConfig(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    ))

    extracted = extract_multiple_sources(sources, chunker)

    return [
        {
            "content": doc.content,
            "metadata": doc.metadata
        }
        for doc in extracted
    ]