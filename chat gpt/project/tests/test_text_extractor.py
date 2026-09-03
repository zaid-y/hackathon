from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.text_extractor import (
    DocumentTextExtractor,
    EmptyDocumentError,
    ExtractionError,
    InvalidDocumentError,
    UnsupportedDocumentError,
)


def test_extracts_utf8_thai_txt_without_inventing_page_number(tmp_path: Path) -> None:
    source = tmp_path / "กติกา.txt"
    source.write_text("วันปิดรับสมัคร 28 กุมภาพันธ์ 2569", encoding="utf-8")

    result = DocumentTextExtractor().extract(source)

    assert result.document == "กติกา.txt"
    assert len(result.pages) == 1
    assert result.pages[0].page is None
    assert "28 กุมภาพันธ์ 2569" in result.pages[0].text


def test_extracts_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    source = tmp_path / "requirements.docx"
    document = DocxDocument()
    document.add_paragraph("คุณสมบัติผู้สมัคร")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "อายุขั้นต่ำ"
    table.cell(0, 1).text = "18 ปี"
    document.save(source)

    result = DocumentTextExtractor().extract(source)

    assert result.pages[0].page is None
    assert "คุณสมบัติผู้สมัคร" in result.pages[0].text
    assert "อายุขั้นต่ำ | 18 ปี" in result.pages[0].text


def test_rejects_empty_txt(tmp_path: Path) -> None:
    source = tmp_path / "empty.txt"
    source.write_text("   \n", encoding="utf-8")

    with pytest.raises(EmptyDocumentError):
        DocumentTextExtractor().extract(source)


def test_reports_malformed_pdf(tmp_path: Path) -> None:
    source = tmp_path / "broken.pdf"
    source.write_bytes(b"this is not a pdf")

    with pytest.raises(InvalidDocumentError):
        DocumentTextExtractor().extract(source)


def test_rejects_unsupported_type(tmp_path: Path) -> None:
    source = tmp_path / "notes.csv"
    source.write_text("question,answer", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentError):
        DocumentTextExtractor().extract(source)


def test_reports_missing_file(tmp_path: Path) -> None:
    with pytest.raises(ExtractionError):
        DocumentTextExtractor().extract(tmp_path / "missing.txt")
